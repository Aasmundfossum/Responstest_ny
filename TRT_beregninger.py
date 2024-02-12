import numpy as np
import pandas as pd
import streamlit as st
import datetime as datetime
from datetime import datetime as dt
import plotly.express as px
import csv
from scipy.stats import linregress
from docx import Document
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
import io
import json
import folium
from streamlit_folium import st_folium
import plotly.graph_objects as go
from html2image import Html2Image
import os

st.set_page_config(page_title="TRT-beregning", page_icon="üî•")

with open("styles/main.css") as f:
    st.markdown("<style>{}</style>".format(f.read()), unsafe_allow_html=True)

class TRT_beregning:
    def __init__(self):
        pass
    
    def kjor_hele(self):
        self.streamlit_input()
        if self.datafil: #and (self.knapp == True):
            self.varmeegenskaper()
            self.les_av_datafil()
            self.finn_alle_tester()
            self.velg_test()
            self.input_til_rapport()
            self.kjor_knapp()
            if self.knapp2 == True:
                with st.spinner(text="Grubler ... ü§î"):
                    self.behandle_test()
                    self.linear_tiln()
                    self.plot1()
                    self.effektiv_varmeledningsevne()
                    self.plot2()
                    self.teoretiske_tempforlop()
                    self.plot3()
                    self.plot4()
                    self.plot5()
                    self.tilfort_effekt()
                    self.plot6()
                    if self.lag_rapport_knapp == True:
                        with st.spinner(text="Grubler ... ü§î"):
                            self.lagre_som_png()
                            self.lag_rapport()

    def streamlit_input(self):
        # Viser alle input-felt i streamlit
        st.title('Resultatberegning fra termisk responstest (TRT) ü•µü•∂')
        st.markdown('Laget av √Ösmund Fossum üë®üèº‚Äçüíª')

        st.header('Last opp fil fra database for automatisk utfylling')
        self.auto = st.file_uploader(label='Last opp JSON-fil fra database for automatisk utfylling (valgfritt)', type='json')
        if self.auto:
            self.inputfil = json.load(self.auto)
        st.markdown('---')

        st.header('Inndata til beregninger') 
        kollvaeske_options = ['HX24','HX35','Kilfrost Geo 24 %','Kilfrost Geo 32 %','Kilfrost Geo 35 %']
        vis_tetthet_input = False
        if self.auto:
            if self.inputfil['Kollektorv√¶ske'] in kollvaeske_options:
                kollvaeske_indeks = kollvaeske_options.index(self.inputfil['Kollektorv√¶ske'])
            else:
                kollvaeske_options.append(self.inputfil['Kollektorv√¶ske'])
                kollvaeske_indeks = kollvaeske_options.index(self.inputfil['Kollektorv√¶ske'])
                vis_tetthet_input = True
        else:
            kollvaeske_indeks = 0
        
        c1, c2 = st.columns(2)
        with c1:
            self.kollvaeske = st.selectbox(label='Type kollektorv√¶ske',options=kollvaeske_options,index=kollvaeske_indeks)
        
        if vis_tetthet_input == True:
            st.write('NB: Fyll inn tettet og varmekapasitet for kollektorv√¶sken!')
            c1, c2, c3, c4 = st.columns([0.1,1,1,0.1])
            with c2:
                self.tetthet = st.number_input('Tetthet til kollektorv√¶ske (kg/m$^3$)')
            with c3:
                self.varmekap = st.number_input('Varmekapasitet til kollektorv√¶ske (kJ/kg K)')
            st.markdown('---')
        

        if self.auto:
            dybde_verdi = self.inputfil['Kollektorlengde']
            diam_verdi = self.inputfil['Br√∏nndiameter']
            strom_foer_verdi = self.inputfil['Str√∏mm√•ler f√∏r']
            strom_etter_verdi = self.inputfil['Str√∏mm√•ler etter']
        else:
            dybde_verdi = 0
            diam_verdi = 0
            strom_foer_verdi = 0
            strom_etter_verdi = 0
        
        c1, c2 = st.columns(2)
        with c1:
            self.dybde = st.number_input(label='Br√∏nndybde (m)', value=dybde_verdi, step=1)
        with c2:
            self.diam = st.number_input(label='Diameter (mm)', value=diam_verdi, step=1)

        c1, c2 = st.columns(2)
        with c1:
            self.strom_foer = st.number_input(label='Str√∏mm√•ler f√∏r test (kWh)', value=strom_foer_verdi, step=1)
        with c2:
            self.strom_etter = st.number_input(label='Str√∏mm√•ler etter test (kWh)', value=strom_etter_verdi, step=1)

        self.datafil = st.file_uploader(label='Datafil fra testrigg (.dat)',type='dat')

        st.markdown('---')

    def varmeegenskaper(self):
        if self.kollvaeske == 'HXi24':
            self.tetthet = 970.5
            self.varmekap = 4.298
        elif self.kollvaeske == 'HXi35':
            self.tetthet = 955
            self.varmekap = 4.061
        elif self.kollvaeske == 'Kilfrost Geo 24 %':
            self.tetthet = 1105.5
            self.varmekap = 3.455 #kJ/kg K
        elif self.kollvaeske == 'Kilfrost Geo 32 %':
            self.tetthet = 1136.2
            self.varmekap = 3.251 #kJ/kg K
        elif self.kollvaeske == 'Kilfrost Geo 35 %':
            self.tetthet = 1150.6
            self.varmekap = 3.156 #kJ/kg K

    def les_av_datafil(self):
        @st.cache_data
        def funk_les_datafil(filnavn):
            df = pd.read_csv(filnavn, 
                    sep=",", 
                    skiprows=1, 
                    usecols=[0,6,7,8,9,10,11,12,13,23],
                    header=0,
                    names = ['tidspunkt','temp_fra_bronn','aPump','temp_til_bronn','rigg','ute','sirk_hast','pumpehast','panelhast','avpaa'])
            
            df = df.drop([0,1])
            df['temp_fra_bronn'] = df['temp_fra_bronn'].astype(float)
            df['aPump'] = df['aPump'].astype(float)
            df['temp_til_bronn'] = df['temp_til_bronn'].astype(float)
            df['rigg'] = df['rigg'].astype(float)
            df['ute'] = df['ute'].astype(float)
            df['sirk_hast'] = df['sirk_hast'].astype(float)
            df['pumpehast'] = df['pumpehast'].astype(float)
            df['panelhast'] = df['panelhast'].astype(float)
            df['avpaa'] = df['avpaa'].astype(int)
            return df
        
        self.df = funk_les_datafil(self.datafil)

    def finn_alle_tester(self):
        def custom_to_datetime(timestamp):
            try:
                return pd.to_datetime(timestamp, format="%Y-%m-%d %H:%M:%S.%f")
            except ValueError:
                return pd.to_datetime(timestamp, format="%Y-%m-%d %H:%M:%S")
            
        @st.cache_data
        def funk_finn_alle_tester(df):
        
            df['tidspunkt'] = df['tidspunkt'].apply(custom_to_datetime)

            liste_over_alle_tester = []
            i = 0

            while i < len(df):
                start_date = df['tidspunkt'].iloc[i]
                
                for j in range(i + 1, len(df)):
                    difference = (df['tidspunkt'].iloc[j] - df['tidspunkt'].iloc[j - 1]).total_seconds()
                    if difference > 60 or j == len(df)-1:
                        end_date = df['tidspunkt'].iloc[j - 1]
                        break

                en_test = df.loc[(self.df['tidspunkt'] >= start_date) & (df['tidspunkt'] <= end_date)]
                en_test = en_test.reset_index(drop=True)
                
                liste_over_alle_tester.append(en_test)
                
                i = j + 1

            # Henter ut kun de testene som har en viss lengde p√• datasettet
            liste_over_tester = []
            varighet_timer = []
            varighet_minutt = []
            for i in range(0,len(liste_over_alle_tester)):
                if len(liste_over_alle_tester[i]) > 20: # Hvis det er en reell test
                    startdiff = (liste_over_alle_tester[i].iloc[9,0]-liste_over_alle_tester[i].iloc[8,0]).total_seconds()
                    sluttdiff = (liste_over_alle_tester[i].iloc[-8,0]-liste_over_alle_tester[i].iloc[-9,0]).total_seconds()
                    varighet = (liste_over_alle_tester[i].iloc[-1,0]-liste_over_alle_tester[i].iloc[0,0]).total_seconds()
            
                    if (startdiff == 5.0 and sluttdiff == 5.0 and varighet >= 180) or (startdiff == 30.0 and varighet >= 72000) or (startdiff == 5.0 and sluttdiff == 30.0 and varighet > 72000): #Hvis uforst_temp m√•ling varer lengre enn 3 minutter eller hoveddel varer lengre enn 20 timer   
                        liste_over_tester.append(liste_over_alle_tester[i])
                        varighet_timer.append(varighet//3600)
                        varighet_minutt.append((varighet-(varighet//3600)*3600)/60)

            return liste_over_tester,varighet_timer,varighet_minutt
        
        self.liste_over_tester,self.varighet_timer,self.varighet_minutt = funk_finn_alle_tester(self.df)



    def velg_test(self):
        self.vis_kjor_knapp = False
        if len(self.liste_over_tester)<1:
            st.markdown('Feil: Finner ingen gyldige tester i denne filen.')
        else:
            liste_over_tester_st = []
            liste_over_tester_st_kun_hoved = []
            liste_over_tester_st_kun_uforst = []

            for i in range(0,len(self.liste_over_tester)):
                startdiff = (self.liste_over_tester[i].iloc[9,0]-self.liste_over_tester[i].iloc[8,0]).total_seconds()
                sluttdiff = (self.liste_over_tester[i].iloc[-8,0]-self.liste_over_tester[i].iloc[-9,0]).total_seconds()
                if startdiff != 5.0:
                    self.type_test = 'Kun hoveddel'
                    tekst_kun_hoved = f'Hoveddel:  ({self.liste_over_tester[i].iloc[0,0].strftime("%d.%m.%Y")}  -  {self.liste_over_tester[i].iloc[-1,0].strftime("%d.%m.%Y")})'
                    liste_over_tester_st_kun_hoved.append(tekst_kun_hoved)
                elif sluttdiff != 30.0:
                    self.type_test = 'Kun m√•ling av uforstyrret temperatur'
                    tekst_kun_uforst = f'M√•ling av uforstyrret temperatur:  ({self.liste_over_tester[i].iloc[0,0].strftime("%d.%m.%Y")}  -  {self.liste_over_tester[i].iloc[-1,0].strftime("%d.%m.%Y")})'
                    liste_over_tester_st_kun_uforst.append(tekst_kun_uforst)
                else:
                    self.type_test = 'Komplett'
                
                tekst = f'Test nr. {i+1}:  ({self.liste_over_tester[i].iloc[0,0].strftime("%d.%m.%Y")}  -  {self.liste_over_tester[i].iloc[-1,0].strftime("%d.%m.%Y")}), ({round(self.varighet_timer[i])} h, {round(self.varighet_minutt[i])} min.) - {self.type_test}'
                liste_over_tester_st.append(tekst)

        
        
            siste_indeks = len(liste_over_tester_st)-1
            valg_av_test = st.selectbox('Velg aktuell test', options = liste_over_tester_st, index = siste_indeks)
            st.markdown('---')
            

            if 'Kun m√•ling av uforstyrret temperatur' in valg_av_test:
                test_del1 = self.liste_over_tester[liste_over_tester_st.index(valg_av_test)]
                self.test_del1 = test_del1
                
                kun_uforst_knapp = st.checkbox('Vis resultater for kun m√•ling av uforstyrret temperatur')

                if kun_uforst_knapp == True:
                    self.type_test = 'Kun m√•ling av uforstyrret temperatur'
                    self.vis_kjor_knapp = True
                else:
                    st.markdown('ELLER:')
                    valg_test_del2 = st.selectbox('Velg tilh√∏rende hoveddel', options = liste_over_tester_st, index=0)
                    if 'Kun hoveddel' in valg_test_del2:
                        self.vis_kjor_knapp = True
                        test_del2 = self.liste_over_tester[liste_over_tester_st.index(valg_test_del2)]
                        self.test_del2 = test_del2
                    else:
                        st.markdown('---')
                        st.markdown('MERK: Nederst m√• det velges en hoveddel. Velg eventuelt √• vise resultater for kun m√•ling av uforstyrret temperatur.')
                
            elif 'Kun hoveddel' in valg_av_test:
                test_del2 = self.liste_over_tester[liste_over_tester_st.index(valg_av_test)]
                self.test_del2 = test_del2

                kun_hoved_knapp = st.checkbox('Bruk kun hoveddel')
                
                if kun_hoved_knapp == False:
                    st.markdown('ELLER:')
                    
                    liste_over_tester_st_steg2 = liste_over_tester_st
                    #for i in range(0,len(liste_over_tester_st_steg2)):
                    #    if 'Kun m√•ling av uforstyrret temperatur' not in liste_over_tester_st_steg2[i]:
                    #        liste_over_tester_st_steg2[i]='-'
                    
                    valg_test_del1 = st.selectbox('Velg tilh√∏rende m√•ling av uforstyrret temperatur',options = liste_over_tester_st_steg2, index=siste_indeks)
                    
                    if 'Kun m√•ling av uforstyrret temperatur' in valg_test_del1:
                        self.vis_kjor_knapp = True
                        
                        test_del1 = self.liste_over_tester[liste_over_tester_st.index(valg_test_del1)]
                        self.test_del1 = test_del1
                    else:
                        st.markdown('---')
                        st.markdown('MERK: Nederst m√• det velges en m√•ling for uforstyrret temperatur. Velg eventuelt √• vise resultater for kun hoveddel.')
                    
                
                elif kun_hoved_knapp == True:
                    self.konst_uforst_temp = st.number_input('Uforstyrret temperatur',value=7.5)
                    self.type_test = 'Kun hoveddel'
                    self.vis_kjor_knapp = True

            elif 'Komplett' in valg_av_test:
                denne_test = self.liste_over_tester[liste_over_tester_st.index(valg_av_test)]
                for i in range(0,len(denne_test)):
                    if denne_test.iloc[i,0] == denne_test.iloc[i-1,0]:
                        deleindeks = int(i)
                        break
                self.vis_kjor_knapp = True
                test_del1 = denne_test.iloc[0:deleindeks,:]
                self.test_del1 = test_del1
                test_del2 = denne_test.iloc[deleindeks:,:]
                self.test_del2 = test_del2

    def input_til_rapport(self):
        st.header('Inndata til Rapport')
        if self.auto:
            sted_verdi = self.inputfil['Prosjektnavn']
            oppdragsgiver_verdi = '' #oppdragsgiver_verdi = self.inputfil['Oppdragsgiver']
        else:
            sted_verdi = ''
            oppdragsgiver_verdi = ''


        c1, c2 = st.columns(2)
        with c1:
            self.sted = st.text_input('Sted hvor responstesten er gjennomf√∏rt', value=sted_verdi)
        with c2:
            self.oppdragsgiver = st.text_input('Oppdragsgiver', value=oppdragsgiver_verdi)
        
        st.markdown('---')
        if self.auto:
            st.write('Kart som viser plassering av br√∏nnen slik det blir i rapporten')
            self.kart = folium.Map(location=[self.inputfil['Latitude'], self.inputfil['Longitude']], zoom_start=16)
            folium.CircleMarker([self.inputfil['Latitude'], self.inputfil['Longitude']], fill=True, color='red', fill_opacity=1, radius=5).add_to(self.kart)
            st_folium(self.kart,height=450,width=800)
        else:
            st.write('Sett inn funksjon for √• skrive inn koordinater og/eller klikke p√• kart')
#            if 'lat' not in st.session_state:
#                st.session_state.lat = 0.0
#            if 'long' not in st.session_state:
#                st.session_state.long = 0.0
#                        
#            st.write('Klikk p√• kartet for √• velge plassering av br√∏nnen:')
#            m = folium.Map(location=[63, 10], zoom_start=16)
#            if st.session_state.lat != None:
#                folium.CircleMarker([st.session_state.lat, st.session_state.long], fill=True, color='red', fill_opacity=1, radius=5).add_to(m)
#            coord_json = st_folium(m,height=500,width=725)
#            
#            if coord_json['last_clicked']:
#                lat = coord_json['last_clicked']['lat']
#                long = coord_json['last_clicked']['lng']
#                st.session_state.lat = float(lat)
#                st.session_state.long = float(long)
        st.markdown('---')
        st.write('Temperaturm√•linger i br√∏nn f√∏r og etter test (Figuren inkluderes i rapporten)')
        
        def json_list_to_python_list(str_fra_json):
            liste = str_fra_json.strip('[]').split()
            for i in range(len(liste)):
                try:
                    if liste[i].lower() == 'none':
                        liste[i] = None
                    else:
                        liste[i] = float(liste[i])
                except ValueError:
                    pass  # Leave it as None if it cannot be converted
            return liste

        if self.auto:
                posisjon_liste_foer = json_list_to_python_list(self.inputfil["Posisjoner temperaturm√•linger f√∏r test"])
                temp_liste_foer = json_list_to_python_list(self.inputfil["Temperaturm√•linger f√∏r test"])
                posisjon_liste_etter = json_list_to_python_list(self.inputfil["Posisjoner temperaturm√•linger etter test"])
                temp_liste_etter = json_list_to_python_list(self.inputfil["Temperaturm√•linger etter test"])

                dato_foer = dt.strptime(self.inputfil['M√•ledato temperaturprofil f√∏r test'], "%Y-%m-%d")
                dato_etter = dt.strptime(self.inputfil['M√•ledato temperaturprofil etter test'], "%Y-%m-%d")

                min_x = np.min([x for x in temp_liste_foer if x is not None])-0.5
                max_x = np.max([x for x in temp_liste_etter if x is not None])+0.5
                snittemp_foer = np.mean([x for x in temp_liste_foer if x is not None])
                grunnvann_foer = self.inputfil['Grunnvannsniv√• f√∏r test']

                fig = px.line()
        
                fig.add_trace(go.Scatter(y=posisjon_liste_foer, x=temp_liste_foer, mode='lines+markers', name=f"F√∏r test; {dato_foer.strftime('%d.%m.%Y')}", line=dict(color='#367A2F')))
                fig.add_trace(go.Scatter(x=[snittemp_foer, snittemp_foer], y=[min(posisjon_liste_foer), max(posisjon_liste_foer)], mode='lines',
                         name='Gjennomsnitt', line=dict(color='#367A2F', width=2, dash='dash')))
                fig.add_trace(go.Scatter(y=posisjon_liste_etter, x=temp_liste_etter, mode='lines+markers', name=f"Etter test; {dato_etter.strftime('%d.%m.%Y')}",
                         line=dict(color='#FFC358')))
                fig.add_trace(go.Scatter(x=[min(temp_liste_foer)*0.1, max(temp_liste_foer)*10], y=[grunnvann_foer, grunnvann_foer], mode='lines',
                         name='Grunnvansniv√• f√∏r test', line=dict(color='blue', width=2, dash='dash')))
                
                fig.update_layout(legend=dict(orientation='h', yanchor='top', y=1.15, xanchor='right', x=1))
                fig.update_xaxes(title_text='Temperatur [¬∞C]', side='top')
                fig.update_yaxes(title_text='Dybde [m]', autorange='reversed')
                fig.update_xaxes(range=[min_x, max_x])
                #fig.update_layout(showlegend=True)
                
                #fig.update_layout(height=750)  # Set the height in pixels
                st.plotly_chart(fig, config={'staticPlot': True}, use_container_width=True)
                self.fig0 = fig
        else:
            st.write('Sett inn funksjonen fra input-appen som lar deg skrive inn temperaturm√•linger manuelt')
        
        st.markdown('---')

    def kjor_knapp(self):
        if self.vis_kjor_knapp == True and self.dybde > 0 and self.diam > 0 and self.strom_etter >0:
            c1,c2,c3 = st.columns(3)
            with c2:
                self.knapp2 = st.checkbox('Kj√∏r beregninger! :steam_locomotive:')
        else:
            st.write('Fyll inn all n√∏dvendig informasjon for √• aktivere kj√∏r-knapp')
            self.knapp2 = False

    def behandle_test(self):
        # Henter ut rader kun for de aktuelle tidspunkter basert p√• hvor "Pump enabeled"-kolonnen sl√•r inn:
        if self.type_test != 'Kun hoveddel':
            for i in range(0,len(self.test_del1)):
                if self.test_del1['avpaa'].iloc[i] != 0:
                    startrad = int(i)
                    break
            self.test_del1 = self.test_del1.iloc[startrad:]

            self.test_del1 = self.test_del1.reset_index(drop=True)                # Resetter radnummer til Dataframes
            # Regner ut gjennomsnittstemperaturen til kollektorv√¶sken, og legger dette som en egen kolonne:
            snittemp = pd.DataFrame((self.test_del1['temp_fra_bronn']+self.test_del1['temp_til_bronn'])/2)
            self.test_del1.insert(4,"snittemp",snittemp,True)
            self.test_del1 = self.test_del1

        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            self.test_del2 = self.test_del2.reset_index(drop=True)
            # Regner ut gjennomsnittstemperaturen til kollektorv√¶sken, og legger dette som en egen kolonne:
            snittemp = pd.DataFrame((self.test_del2['temp_fra_bronn']+self.test_del2['temp_til_bronn'])/2)
            self.test_del2.insert(4,"snittemp",snittemp,True)

            # Lager kolonne som viser antall sekunder siden testen startet
            startindeks = 0            #Raden hvor testen faktisk starter
            antall_sek = pd.DataFrame({'sek_siden_start' : [0]*len(self.test_del2)})
            for i in range(0,len(self.test_del2)):
                antall_sek.iloc[i] = (self.test_del2['tidspunkt'].iloc[i]-self.test_del2['tidspunkt'].iloc[startindeks]).total_seconds()

            self.test_del2.insert(1,"sek_siden_start",antall_sek,True)                  # Setter inn antall_sek som kolonne (indeks 1) i test_del2-dataframen
            self.test_del2['sek_siden_start'] = self.test_del2['sek_siden_start'].astype(float)

            # Definerer ln(t)
            ln_t = pd.DataFrame({'ln_t' : [0]*len(self.test_del2)})
            for i in range(0,len(self.test_del2)):
                ln_t.iloc[i] = np.log(self.test_del2['sek_siden_start'].iloc[i])
            
            self.test_del2.insert(2,"ln_t",ln_t,True)
            self.test_del2['ln_t'] = self.test_del2['ln_t'].astype(float)

            # Henter ut den delen av test_del2 som foreg√•r etter 5 og 20 timer:
            etter5timer = self.test_del2.iloc[600:,:]                  #Antar her at det er et 30 sek mellom hvert m√•lepunkt
            self.etter5timer = etter5timer.reset_index(drop=True)
            etter20timer = self.test_del2.iloc[2400:,:]                  #Antar her at det er et 30 sek mellom hvert m√•lepunkt
            self.etter20timer = etter20timer.reset_index(drop=True)
            self.test_del2 = self.test_del2


    def linear_tiln(self):
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            # Line√¶r tiln√¶rming av gjennomsnittstemperaturen etter 20 timer:

            slope, intercept, r_value, p_value, std_err = linregress(self.etter20timer['ln_t'], self.etter20timer['snittemp'])
            self.y_pred = slope * self.etter20timer['ln_t'] + intercept
            self.r_verdi = str(round(r_value,3))

    def plot1(self):
        st.markdown('---')
        st.header('Resultater fra beregninger')

        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            til_plot1 = pd.DataFrame({"Tider" : self.etter20timer['ln_t'], "Gj.snittstemp. kollektorv√¶ske" : self.etter20timer['snittemp'], "Line√¶r tiln√¶rming med r = "+self.r_verdi+"" : self.y_pred})
            fig1 = px.line(til_plot1, x='Tider', y=["Gj.snittstemp. kollektorv√¶ske","Line√¶r tiln√¶rming med r = "+self.r_verdi+""], title='Temperaturm√•linger etter 20 timer', color_discrete_sequence=['#367A2F', '#FFC358'])
            fig1.update_layout(xaxis_title='Logaritmen av tid siden teststart', yaxis_title='Temperatur (\u2103)',legend_title=None)
            fig1.update_xaxes(range=[11.0, 12.6], row=1, col=1)
            st.plotly_chart(fig1)
            self.fig1 = fig1
 
    def effektiv_varmeledningsevne(self):
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            # Utregning av effektiv varmeledningsevne:
            self.indeks5timer = 600          # Dersom det er 30 sek mellom hver m√•ling
            hjelpefunk = np.zeros(len(self.test_del2))
    
            for i in range(self.indeks5timer+1,len(self.test_del2)):
                y_values = np.array(self.test_del2['snittemp'].iloc[self.indeks5timer:i])
                x_values = np.array(self.test_del2['ln_t'].iloc[self.indeks5timer:i])
                hjelpefunk[i], intercept = np.polyfit(x_values, y_values, 1)
            
            tot_varighet_timer = self.test_del2['sek_siden_start'].iloc[-1]/3600
            mid_effekt = (self.strom_etter-self.strom_foer)/(tot_varighet_timer)*1000    # W
            
            effekt_per_m = (mid_effekt/self.dybde) # W
            ledn_evne = effekt_per_m/(4*np.pi*hjelpefunk)
            for i in range(0,len(ledn_evne)):                   # Fjerner urimelig h√∏ye verdier
                if ledn_evne[i] >= 20:
                    ledn_evne[i] = float('nan')
            ledn_evne = pd.DataFrame(ledn_evne)

            self.stabil_ledn_evne = float(ledn_evne.tail(self.indeks5timer).mean())      # Antar at ledningsevnen stabiliserer seg p√• gjennomsnittet av de siste 5 timer

            self.mid_effekt = mid_effekt
            self.ledn_evne = ledn_evne

    def plot2(self):
        st.markdown("---")
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            #self.ledn_evne_slider = st.slider('Varmeledningsevne som kurven konvergerer mot', 0.1, float(10), self.stabil_ledn_evne, step=0.01)
            self.ledn_evne_slider = st.number_input('Varmeledningsevne som kurven konvergerer mot. Velg verdien som passer best til plottet under.', 0.1*self.stabil_ledn_evne, 2*self.stabil_ledn_evne, self.stabil_ledn_evne, step=0.01)
            stabil_ledn_evne_tilplot = pd.DataFrame({'Value' : [self.ledn_evne_slider]*len(self.test_del2)})
            til_plot2 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start'].iloc[self.indeks5timer:]/3600, "Effektiv ledningsevne" : self.ledn_evne[0].iloc[self.indeks5timer:], 'Stabilisert ledningsevne' : stabil_ledn_evne_tilplot['Value']})
            fig2 = px.line(til_plot2, x='Tider', y=['Effektiv ledningsevne','Stabilisert ledningsevne'], title='Utvikling av effektiv varmeledningsevne', color_discrete_sequence=['#367A2F', '#FFC358'])
            fig2.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Effektiv varmeledningsevne (W/mK)',legend_title=None)
            fig2.update_xaxes(range=[0, self.test_del2['sek_siden_start'].iloc[-1]/3600], row=1, col=1)
            fig2.update_yaxes(range=[self.stabil_ledn_evne*0.5, self.stabil_ledn_evne*1.5], row=1, col=1)
            st.plotly_chart(fig2)
            self.fig2 = fig2
        
    def teoretiske_tempforlop(self):    
        if self.type_test != 'Kun hoveddel':
            # Regner ut teoretiske temperaturforl√∏p med gitte verdier av borehullsmotstand:
            snitt_temp_inn = self.test_del1.loc[10:, 'temp_til_bronn'].mean()
            snitt_temp_ut = self.test_del1.loc[10:, 'temp_fra_bronn'].mean()                      #Bruker alle temp. m√•linger unntatt de 10 f√∏rste
            self.uforst_temp_verdi = float((snitt_temp_inn+snitt_temp_ut)/2)
        else:
            self.uforst_temp_verdi = self.konst_uforst_temp

        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            lambdaa = self.ledn_evne_slider             # Den verdien som kurven i fig 2 stabiliserer seg mot.
            varmekap_Jm3K = 2200000
            self.diff = lambdaa/varmekap_Jm3K 
            #self.motstand = st.slider('Varmemotstand (mK/W)', float(0), float(1), 0.08, step=0.01)
            self.I1 = (self.mid_effekt/(4*np.pi*lambdaa*self.dybde))
            motstand_gjett_vektor = (np.array(self.test_del2['snittemp'])-self.uforst_temp_verdi-self.I1*np.array(self.test_del2['ln_t'])-self.I1*(np.log((4*self.diff)/(((self.diam/1000)/2)**2))-0.5772)) * (self.dybde/self.mid_effekt)
            self.motstand_gjett = np.mean(motstand_gjett_vektor[3:])

            st.markdown("---")

    def plot3(self):
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            self.motstand = st.number_input('Varmemotstand (mK/W). Velg den verdien som gir best samsvar med kurven under.', float(0), 1.5*self.motstand_gjett, self.motstand_gjett, step=0.0001)
            teori_temp = self.uforst_temp_verdi+self.I1*np.array(self.test_del2['ln_t'])+self.I1*(np.log((4*self.diff)/(((self.diam/1000)/2)**2))-0.5772)+(self.mid_effekt/self.dybde)*self.motstand
            self.teori_temp = pd.DataFrame(teori_temp)

            self.motstand_str = str(round(self.motstand,3))
            snittemp_tilplot = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Gj.snittstemp. kollektorv√¶ske" : self.test_del2['snittemp'], "Typekurve R = "+self.motstand_str+" mK/W" : self.teori_temp[0]})
            fig3 = px.line(snittemp_tilplot, x='Tider', y=["Gj.snittstemp. kollektorv√¶ske","Typekurve R = "+self.motstand_str+" mK/W"], title='Faktisk temp.forl√∏p og typekurve for termisk motstand R = '+self.motstand_str+' mK/W', color_discrete_sequence=['#367A2F', '#FFC358'])
            fig3.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Temperatur (\u2103)',legend_title=None)
            st.plotly_chart(fig3)
            self.fig3 = fig3

    def plot4(self):
        st.markdown("---")
        
        if self.type_test != 'Kun hoveddel':
            # Definerer uforstyrret temperatur:
            uforst_temp = pd.DataFrame({'Value' : [self.uforst_temp_verdi]*len(self.test_del1)})
            df_tilplot = pd.DataFrame({"Tider" : self.test_del1['tidspunkt'], "Gj.snittstemp. kollektorv√¶ske" : self.test_del1['temp_fra_bronn'], "Uforstyrret temperatur" : uforst_temp['Value']})
            fig4 = px.line(df_tilplot, x='Tider', y=["Gj.snittstemp. kollektorv√¶ske","Uforstyrret temperatur"], title='Temperaturm√•linger fra sirkulasjon av kollektorv√¶ske', color_discrete_sequence=['#367A2F', '#FFC358'])
            fig4.update_layout(xaxis_title='Tid (klokkeslett)', yaxis_title='Temperatur (\u2103)',legend_title=None)
            st.plotly_chart(fig4)
            self.fig4 = fig4

    def plot5(self):
        st.markdown("---")
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            til_plot5 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Temperatur til br√∏nnen" : self.test_del2['temp_til_bronn'], "Temperatur fra br√∏nnen" : self.test_del2['temp_fra_bronn'], "Temperatur i testrigg" : self.test_del2['rigg'], "Temperatur i uteluft" : self.test_del2['ute']})
            fig5 = px.line(til_plot5, x='Tider', y=['Temperatur til br√∏nnen', 'Temperatur fra br√∏nnen', 'Temperatur i testrigg', 'Temperatur i uteluft'], title='Utelufttemp. og kollektorv√¶sketemp. til og fra energibr√∏nnen', color_discrete_sequence=['#367A2F', '#C2CF9F', '#FFC358', '#FFE7BC'])
            fig5.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Temperatur (\u2103)',legend_title=None)
            st.plotly_chart(fig5)
            self.fig5 = fig5

    def tilfort_effekt(self):
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            # Regner ut tilf√∏rt varmeeffekt ved bruk av sirkulasjonshastighet:
            tilfort_eff = self.tetthet*(np.array(self.test_del2['sirk_hast'])/60)/1000*self.varmekap*np.abs(np.array(self.test_del2['temp_til_bronn'])-np.array(self.test_del2['temp_fra_bronn']))
            self.tilfort_eff = pd.DataFrame(tilfort_eff)

    def plot6(self):
        st.markdown("---")
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            til_plot6 = pd.DataFrame({"Tider" : self.test_del2['sek_siden_start']/3600, "Tilf√∏rt varmeeffekt" : self.tilfort_eff[0], "Sirkulasjonshastighet" : self.test_del2['sirk_hast']})
            fig6 = px.scatter(til_plot6, x='Tider', y=["Tilf√∏rt varmeeffekt","Sirkulasjonshastighet"], title='Sirkulasjonshastighet og tilf√∏rt varmeeffekt', color_discrete_sequence=['#367A2F', '#FFC358'])
            fig6.update_layout(xaxis_title='Tid siden teststart (timer)', yaxis_title='Sirkulasjonsmengde [l/min] og tilf√∏rt effekt [kW]',legend_title=None)
            st.plotly_chart(fig6)
            self.fig6 = fig6

        st.markdown("---")
        self.lag_rapport_knapp = st.checkbox('Generer rapport üìù')

    def lagre_som_png(self):
        fig_bredde = 800
        fig_hoyde = 450

        self.kart.save("TRT-figurer/fig_kart.html")
        hti = Html2Image(custom_flags=['--virtual-time-budget=10000', '--hide-scrollbars'], output_path='TRT-figurer')

        if os.path.exists("TRT-figurer/fig_kart.png"):
            os.remove("TRT-figurer/fig_kart.png")  # Remove the existing file
        
        hti.screenshot(html_file="TRT-figurer/fig_kart.html", save_as="fig_kart.png", size=(fig_bredde,fig_hoyde))
        self.fig0.write_image("TRT-figurer/fig0.png",width=fig_bredde, height=fig_hoyde)
        
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            self.fig1.write_image("TRT-figurer/fig1.png",width=fig_bredde, height=fig_hoyde)
            self.fig2.write_image("TRT-figurer/fig2.png",width=fig_bredde, height=fig_hoyde)
            self.fig3.write_image("TRT-figurer/fig3.png",width=fig_bredde, height=fig_hoyde)
            self.fig5.write_image("TRT-figurer/fig5.png",width=fig_bredde, height=fig_hoyde)
            self.fig6.write_image("TRT-figurer/fig6.png",width=fig_bredde, height=fig_hoyde)
        elif self.type_test != 'Kun hoveddel':
            self.fig4.write_image("TRT-figurer/fig4.png",width=fig_bredde, height=fig_hoyde)

    def lag_rapport(self):        
        document = Document("Mal Rapport TRT - Python.docx")
        styles = document.styles
        style = styles.add_style('Citation', WD_STYLE_TYPE.PARAGRAPH)
        
        # Setter inn stedsnavn, tall osv i rapporten ved gitte kode-substrings:
        def sett_inn_i_rapport(eks_str,ny_str):
            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if eks_str in paragraph.text:
                    linje_til_motstand = paragraph_index
                    innhold = document.paragraphs[linje_til_motstand].text
                    document.paragraphs[linje_til_motstand].clear()
                    document.paragraphs[linje_til_motstand].text = innhold.replace(eks_str,ny_str)
                    
        sett_inn_i_rapport("[python_sted]", self.sted)
        sett_inn_i_rapport("[python_bronnborer]",self.oppdragsgiver)
        sett_inn_i_rapport("[python_dybde]",str(self.dybde))
        sett_inn_i_rapport("[python_ledn_evne]",str(round(self.ledn_evne_slider,2)))
        sett_inn_i_rapport("[python_motstand]",self.motstand_str)
        sett_inn_i_rapport("[python_uforst_temp]",str(round(self.uforst_temp_verdi, 2)))
        sett_inn_i_rapport("[python_0_02_pluss_motstand]",str(round(self.motstand+0.02, 2)))

        # Setter inn figurer p√• riktig sted i rapporten:
        if self.type_test != 'Kun m√•ling av uforstyrret temperatur':
            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_kart]" in paragraph.text:
                    linje_til_kart = paragraph_index
                    break  
            document.paragraphs[linje_til_kart].clear()
            run_kart = document.paragraphs[linje_til_kart].add_run()
            kart = run_kart.add_picture("TRT-figurer/fig_kart.png")
            kart.width = Cm(16)
            kart.height = Cm(9)

            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_fig0]" in paragraph.text:
                    linje_til_fig0 = paragraph_index
                    break  
            document.paragraphs[linje_til_fig0].clear()
            run0 = document.paragraphs[linje_til_fig0].add_run()
            bilde0 = run0.add_picture("TRT-figurer/fig0.png")
            bilde0.width = Cm(16)
            bilde0.height = Cm(9)
            
            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_fig2]" in paragraph.text:
                    linje_til_fig2 = paragraph_index
                    break  
            document.paragraphs[linje_til_fig2].clear()
            run2 = document.paragraphs[linje_til_fig2].add_run()
            bilde2 = run2.add_picture("TRT-figurer/fig2.png")
            bilde2.width = Cm(16)
            bilde2.height = Cm(9)

            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_fig3]" in paragraph.text:
                    linje_til_fig3 = paragraph_index
                    break
            document.paragraphs[linje_til_fig3].clear()
            run3 = document.paragraphs[linje_til_fig3].add_run()
            bilde3 = run3.add_picture("TRT-figurer/fig3.png")
            bilde3.width = Cm(16)
            bilde3.height = Cm(9)

            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_fig5]" in paragraph.text:
                    linje_til_fig5 = paragraph_index
                    break
            document.paragraphs[linje_til_fig5].clear()
            run5 = document.paragraphs[linje_til_fig5].add_run()
            bilde5 = run5.add_picture("TRT-figurer/fig5.png")
            bilde5.width = Cm(16)
            bilde5.height = Cm(9)

            for paragraph_index, paragraph in enumerate(document.paragraphs):
                if "[python_fig6]" in paragraph.text:
                    linje_til_fig6 = paragraph_index
                    break
            document.paragraphs[linje_til_fig6].clear()
            run6 = document.paragraphs[linje_til_fig6].add_run()
            bilde6 = run6.add_picture("TRT-figurer/fig6.png")
            bilde6.width = Cm(16)
            bilde6.height = Cm(9)

        # Laster ned rapporten vha. download-knapp
        st.markdown("---")
        bio = io.BytesIO()
        document.save(bio)
        if document:
            st.download_button(
                label="Last ned rapport üìù",
                data=bio.getvalue(),
                file_name=f"Termisk responstest - {self.sted}.docx",
                mime="docx")

TRT_beregning().kjor_hele()
